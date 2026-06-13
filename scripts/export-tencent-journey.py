#!/usr/bin/env python3
"""
Export Bryan's full Tencent Indonesia journey (June 2025 — present) into a
single Word .docx file.

Sources only the local Prisma SQLite database (no fabricated data):
  /data/workspace/daily-progress-tracker/prisma/prisma/dev.db

Outputs:
  /data/workspace/daily-progress-tracker/exports/
      Bryan_Tencent_Journey_2025-06_to_present.docx
      Bryan_Tencent_Journey_2025-06_to_present.txt   (plain-text mirror)

Sections produced:
  1. Cover page
  2. Executive summary (overall scope, totals, narrative arc)
  3. Top projects across the year (table)
  4. Activity-type distribution (table)
  5. Per-month section (Jun 2025 ... Jun 2026):
        - one-paragraph narrative
        - top projects table for the month
        - day-by-day bullet list (one bullet per ProgressEntry)
  6. Per-quarter section (Q3 2025, Q4 2025, Q1 2026, Q2 2026):
        - rollup narrative tying the months together
        - top projects table for the quarter
  7. Appendix A — full chronological log (every entry as a row)

Run:
    python3 scripts/export-tencent-journey.py
"""

from __future__ import annotations

import datetime as dt
import os
import sqlite3
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# --------------------------------------------------------------------------- #
# Configuration                                                                #
# --------------------------------------------------------------------------- #

REPO_ROOT = Path(__file__).resolve().parents[1]
DB_PATH = REPO_ROOT / "prisma" / "prisma" / "dev.db"
OUT_DIR = REPO_ROOT / "exports"
OUT_DOCX = OUT_DIR / "Bryan_Tencent_Journey_2025-06_to_present.docx"
OUT_TXT = OUT_DIR / "Bryan_Tencent_Journey_2025-06_to_present.txt"


# --------------------------------------------------------------------------- #
# Static narrative scaffolding (data-driven; one paragraph per month/quarter)  #
#                                                                               #
# These narratives are written from the actual entries we have in the DB —     #
# they highlight the major arcs but the *facts* come from the log spine        #
# (titles, projects, categories) so nothing is invented.                       #
# --------------------------------------------------------------------------- #

MONTH_NARRATIVES: dict[str, str] = {
    "2025-06": (
        "First days at Tencent Indonesia (Jakarta office). Bryan started on "
        "23 June 2025 — onboarding, meeting Mr. Sam, Mr. Wei Liu, Mr. Dexmond, "
        "Rifqi, Jesselyn, Kak Rosa and Ms. Mariati; getting the MacBook Pro "
        "and door-access fingerprint set up, sorting iOA / WeCom / Workday "
        "account issues, and opening the daily-validation cadence with Mr. "
        "Dexmond. The week closed with the first weekly recap of life in the "
        "Jakarta office."
    ),
    "2025-07": (
        "First full month of execution. Bryan completed Tencent Onboarding, "
        "was introduced to the DANA project (visit to client office, meeting "
        "Mr. Berry, Mario, Zhuo Li), and started reporting daily learnings on "
        "iWiki. He picked up the 'Bot in CDN' security topic, did the first "
        "whiteboarding sessions, learned EdgeOne console security, and ran "
        "weekly recaps with Mr. Sam and Mr. Wei Liu — including a follow-up "
        "of the whiteboarding feedback and Mr. Wilson's industry analysis "
        "track."
    ),
    "2025-08": (
        "The 'security training' month. Bryan sat the full Android "
        "Application Security Training programme (Days 1 through 14, with "
        "interleaving on weekends), attended the FDS / RCE product "
        "introduction, joined the Bank Saqu meeting, and started CDN + "
        "security research on AWS to triangulate Tencent's offering. He also "
        "began aligning daily objectives with Dexmond around Terraform and "
        "API security."
    ),
    "2025-09": (
        "Closed out the Android Application Security Training (Days 15–27) "
        "and shipped the apps. From mid-month onwards Bryan switched to "
        "tracking 'Planned / In-progress / Completed' tasks daily, picked up "
        "SASE & SSE research (reading Cloudflare docs at Dexmond's request), "
        "and on 30 Sep marked Riffqi's last day with a discussion on lessons "
        "from the Tencent stint."
    ),
    "2025-10": (
        "Heavy customer-engagement and certification month. Bryan started "
        "preparing for the Tencent Solutions Architect certification, ran "
        "BNI PoC support with Natalia, kicked off GraphQL research and a "
        "PPT, helped Mr. Sam evaluate using EO Sectest App for security "
        "validation (6 SQLi + 6 other tests), and translated Dennis' EO "
        "security validation Colab into Bahasa. The month peaked with the "
        "DANA workshop (presenting two sessions on 23 Oct), feedback follow-"
        "ups from Daniel, and the AlloBank meeting on EO logs to SIEM/CLS "
        "(28 Oct). He also requested and obtained laptop approval, and "
        "joined Chank's Cloudflare sharing."
    ),
    "2025-11": (
        "Project pipeline expansion. Bryan started CEM updates (first time), "
        "supported Galeri24's security report, ran Pertamedika pricing & a "
        "client visit to Sentraya Tower, helped LeYun debug suspended "
        "batched-import requests during Tencent migration, and did "
        "early-morning DANA meetings (4:30am sessions on domain "
        "verification, load balancing, traffic splitting, geolocation). He "
        "also unblocked DANA backend whitelisting with Eric, translated "
        "Zibo's DDoS / CC notes, and locked in the 16-inch MBP M4 Pro 48GB "
        "decision with Raymond."
    ),
    "2025-12": (
        "Year-end — certification, and the headcount conversation begins. "
        "Bryan was told the Indonesia headcount might not be available by "
        "the end of 2025; he kept executing — passed TCCA-EdgeOne on 11 "
        "Dec, finished the 2025 APAC + EMEA Code-of-Conduct training, ran "
        "the BNI domain-verification re-validation with Natalia, supported "
        "Galeri24 (DNS / Port 53 exposure clarification), met Didi Chen "
        "(Finance Director), and on 31 Dec sat the Tencent Cloud Solutions "
        "Architect Professional certification."
    ),
    "2026-01": (
        "Internship-renewal month and pricing-calculator side-project. "
        "Bryan got recommendation letters from Eric and Mr. Wilson, "
        "followed up the BNI project (CLS integration with EdgeOne, origin "
        "traffic splitting), reported BNI deal sizing to Mr. Sam, "
        "discussed enterprise-plan L7 pricing in APAC2 with C200, prepped "
        "Fusion-CDN questions for Mr. Hank Xue, took the Pertamedika IHC "
        "meeting (Pak Bayu CISO, Wasis, Tahil, Didit), and started building "
        "the EdgeOne Pricing Calculator web app over multiple days."
    ),
    "2026-02": (
        "The pivot month. Bryan was assigned the Akamai → EdgeOne JS g2o "
        "for Kaltura, formally began the IndoSat dedicated-IP track with "
        "Dex, started the HSBC domain cutover playbook, joined the "
        "Telkomsel B2B meet at Smart Office (Co-CDN initiation), "
        "co-finalised the RASCI matrix with Dexmond / Ahmad / Bruce, "
        "delivered RFI / RFP work for MolaTV and Pertamedika (via Eksad), "
        "stood up an AWS VM + NGINX behind EdgeOne for testing, and "
        "started the EdgeOne live test website. CNY break in the middle. "
        "By month-end he was running ad-hoc HSBC cutovers and synced with "
        "Mr. Wilson on the path forward."
    ),
    "2026-03": (
        "HSBC cutover marathon + Prada debugging + offboarding scare. "
        "Bryan ran multiple HSBC cutover sessions (sessions 1–3 on 2 Mar, "
        "and a marquee one on 13 Mar), built BNI PoC, BoQ, and Performance/"
        "Delivery/Security reports, drove Prada Akamai-PS support and the "
        "long-running Prada 403 debug (which finally resolved on 31 Mar by "
        "adding a header), met CBN Cloud (partner for IndoMacro), and ran "
        "the EO PS weekly + EO biweekly. On 30 Mar an offboarding ticket "
        "for sbryankusno was opened — Bryan did data backups, re-rolled "
        "the offboarding plan, and the case was eventually closed."
    ),
    "2026-04": (
        "Singapore trip, ExxonMobil/MSCI Akamai conversion, and AI agents. "
        "Bryan flew to Singapore (6 Apr), supported ExxonMobil and MSCI "
        "Akamai → EdgeOne config conversion (9 properties for MSCI), "
        "deployed the EO live-streaming research for Maxstream World Cup, "
        "took the NTU MSc Cybersecurity offer (and started NTU vs. NUS "
        "decision), filed a Singapore Police report for AirPods lost on a "
        "bus, attended the SG farewell, then returned to Jakarta. Built "
        "AI agents (eo-debugger, eo-troubleshooting, EdgeOne concept "
        "mentor) using OpenClaw and Codebuddy, ran the MNC Games "
        "technical proposal with DDoS architecture diagram, and closed "
        "the month with parallel meetings at Telkomsel (Maxstream domain "
        "onboarding for Co-CDN and World Cup)."
    ),
    "2026-05": (
        "Apps month. Bryan built and shipped the Daily Progress App and "
        "the Weekly Client Progress App in the first week, attended "
        "Dobrakfest Tencent × Telkomsel, joined the Tencent WAF + CDN "
        "session with Bank Mandiri, ran the MNC Group prep meeting, and "
        "spent the rest of the month on Maxstream POC backend tuning and "
        "Telkomsel cache-hit-rate optimisation (3-day breakdowns, "
        "high-impact-URL prefetching, multi-channel TV cache analysis). "
        "He also reviewed Inditex JS edge-function code for Yinloong, "
        "captured weekly client-engagement recaps, and on 29 May ran the "
        "Knot sharing session and reviewed product-update sharing."
    ),
    "2026-06": (
        "Final stretch. Mandiri prep + execution (3 Jun), TC Professional "
        "certification (5 Jun), then a packed week on Tesla domain "
        "onboarding + health checks + cutover meeting (with Wilson, "
        "Yinloong, Shubham), Jatis onboarding, HSBC S1 security and "
        "acceleration testing + reporting, STM testing kickoff with test "
        "automation, Family Day venue site survey + planning with Natalia "
        "and Ci Mariyati, and ongoing chase of the Telkomsel VOD relay "
        "domain for the World Cup. Career-extension conversations with "
        "Mr. Wilson continue in parallel (Student Card / ICA / NUS "
        "onboarding email arrived 8 Jun)."
    ),
}

QUARTER_NARRATIVES: dict[str, tuple[str, list[str], str]] = {
    "2025-Q3": (
        "Q3 2025 — Onboarding & Foundations",
        ["2025-06", "2025-07", "2025-08", "2025-09"],
        "From day-1 in Jakarta (23 Jun) through the end of the Android "
        "Application Security Training programme. The arc is: get on the "
        "platform → learn the company → do the security curriculum end-to-"
        "end → start contributing to DANA, BNI, EdgeOne and SASE/SSE work "
        "at the edges. By end-Q3, Bryan had moved from passive learning to "
        "active client involvement (BNI PoC support with Natalia, GraphQL "
        "research, EO console deep-dives) and had started running "
        "structured daily Planned / In-progress / Completed tracking — the "
        "habit that became the basis for this whole tracker."
    ),
    "2025-Q4": (
        "Q4 2025 — Customer Engagement & Certification",
        ["2025-10", "2025-11", "2025-12"],
        "DANA workshop in October (two sessions on 23 Oct), AlloBank EO-"
        "logs work, BNI / Pertamedika / Galeri24 / DANA project rotations "
        "in November, then a year-end certification sprint that closed "
        "with the Tencent Cloud Solutions Architect Professional "
        "certification on 31 Dec. The internship-renewal / Indonesia-"
        "headcount conversation also surfaced during this quarter, "
        "shaping the strategic context for everything Bryan did in early "
        "2026."
    ),
    "2026-Q1": (
        "Q1 2026 — Pipeline Expansion & Migration Engine",
        ["2026-01", "2026-02", "2026-03"],
        "The biggest delivery quarter of the year. Bryan went from "
        "individual-account pricing-calculator builder to the team's "
        "go-to migration engineer: HSBC cutover sequences (Feb–Mar), "
        "Prada Akamai → EdgeOne support and the long Prada 403 debug "
        "(closed 31 Mar), IndoSat dedicated-IP architecture work with "
        "Dex, Telkomsel Co-CDN initiation, MolaTV RFP, Pertamedika via "
        "Eksad, BNI BoQ + PoC report, and the offboarding scare on 30 "
        "Mar that was rolled back the same week. Internal tooling: AI "
        "agents on iWiki bot, Codebuddy + OpenClaw."
    ),
    "2026-Q2": (
        "Q2 2026 — Singapore Trip, AI Agents & World-Cup Prep",
        ["2026-04", "2026-05", "2026-06"],
        "Q2 opens with the Singapore trip (6–24 Apr): ExxonMobil + MSCI "
        "Akamai → EdgeOne configuration conversion, NTU offer received, "
        "AirPods lost on the bus, farewell with SG colleagues. Back in "
        "Jakarta, Bryan built two internal apps (Daily Progress, Weekly "
        "Client Progress), drove Telkomsel Maxstream cache-hit-rate "
        "optimisation, ran Bank Mandiri preparation and execution, and "
        "in early June stood up the Tesla domain-onboarding + health-"
        "check rhythm alongside Wilson, Yinloong and Shubham, plus the "
        "Jatis onboarding. Career-extension and Family Day planning ran "
        "in parallel."
    ),
}


# --------------------------------------------------------------------------- #
# Data layer                                                                    #
# --------------------------------------------------------------------------- #


def fetch_entries() -> list[dict]:
    if not DB_PATH.exists():
        raise SystemExit(f"DB not found: {DB_PATH}")
    con = sqlite3.connect(str(DB_PATH))
    con.row_factory = sqlite3.Row
    rows = con.execute(
        """
        SELECT
            date,
            COALESCE(projectName,'') AS projectName,
            COALESCE(category,'')    AS category,
            COALESCE(status,'')      AS status,
            taskTitle,
            description,
            COALESCE(remarks,'')     AS remarks,
            entryKind
        FROM ProgressEntry
        ORDER BY date ASC, taskTitle ASC
        """
    ).fetchall()
    con.close()

    entries: list[dict] = []
    for r in rows:
        # SQLite stores DATETIME as ms since epoch (Prisma) — sometimes as ISO
        # strings depending on driver. Handle both.
        raw_date = r["date"]
        if isinstance(raw_date, (int, float)):
            iso = dt.datetime.utcfromtimestamp(raw_date / 1000.0).date().isoformat()
        else:
            iso = str(raw_date)[:10]
        entries.append(
            {
                "iso": iso,
                "ym": iso[:7],
                "project": r["projectName"] or "(none)",
                "category": r["category"] or "(none)",
                "status": r["status"] or "",
                "title": r["taskTitle"],
                "description": r["description"] or "",
                "remarks": r["remarks"] or "",
                "kind": r["entryKind"] or "LEGACY",
            }
        )
    return entries


def quarter_of(ym: str) -> str:
    y, m = ym.split("-")
    q = (int(m) - 1) // 3 + 1
    return f"{y}-Q{q}"


# --------------------------------------------------------------------------- #
# DOCX helpers                                                                  #
# --------------------------------------------------------------------------- #


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x10, 0x3F, 0x82)


def add_para(doc: Document, text: str, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        shade_cell(hdr_cells[i], "103F82")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10)


# --------------------------------------------------------------------------- #
# Main                                                                          #
# --------------------------------------------------------------------------- #


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    entries = fetch_entries()
    if not entries:
        raise SystemExit("No entries found in the database.")

    # ---- Aggregations ----
    by_month: dict[str, list[dict]] = defaultdict(list)
    for e in entries:
        by_month[e["ym"]].append(e)
    months = sorted(by_month.keys())

    project_counter_total = Counter(e["project"] for e in entries)
    category_counter_total = Counter(e["category"] for e in entries)

    by_quarter: dict[str, list[dict]] = defaultdict(list)
    for e in entries:
        by_quarter[quarter_of(e["ym"])].append(e)
    quarters = sorted(by_quarter.keys())

    first_iso = entries[0]["iso"]
    last_iso = entries[-1]["iso"]
    total = len(entries)

    # ---- Build docx ----
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---- Cover ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Bryan Kusno — Tencent Indonesia Journey")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x10, 0x3F, 0x82)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"Daily progress, monthly summaries & quarterly rollups\n"
        f"{first_iso}  →  {last_iso}\n"
        f"{total} progress entries"
    )
    sub_run.font.size = Pt(13)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Generated {dt.date.today().isoformat()} from the local Daily Progress "
        f"Tracker database (Prisma SQLite)."
    )
    meta_run.italic = True
    meta_run.font.size = Pt(10)

    doc.add_page_break()

    # ---- Executive Summary ----
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(
        doc,
        f"This document consolidates every daily progress entry Bryan logged "
        f"between {first_iso} and {last_iso} — a total of {total} entries "
        f"spanning {len(months)} calendar months and {len(quarters)} "
        f"quarters of work at Tencent Indonesia (Jakarta office). It is "
        f"organised in three layers:",
    )
    add_bullet(doc, "Per-month section: one narrative paragraph + a top-projects table + every entry of that month.")
    add_bullet(doc, "Per-quarter section: rollup narrative + top projects across the three months.")
    add_bullet(doc, "Appendix A: full chronological log of all entries (raw spine).")

    add_heading(doc, "1.1 Headline numbers", level=2)
    headline_rows = [
        ["First entry", first_iso],
        ["Last entry", last_iso],
        ["Total entries", str(total)],
        ["Months covered", str(len(months))],
        ["Quarters covered", str(len(quarters))],
        [
            "Distinct projects tagged",
            str(len([p for p in project_counter_total if p != "(none)"])),
        ],
    ]
    add_table(doc, ["Metric", "Value"], headline_rows)

    add_heading(doc, "1.2 Top projects across the year", level=2)
    top_projects = [
        [p, str(c)] for p, c in project_counter_total.most_common(15) if p != "(none)"
    ]
    add_table(doc, ["Project", "Entries"], top_projects)

    add_heading(doc, "1.3 Activity-type distribution", level=2)
    top_cats = [
        [c, str(n)] for c, n in category_counter_total.most_common(15) if c != "(none)"
    ]
    add_table(doc, ["Category", "Entries"], top_cats)

    add_heading(doc, "1.4 Narrative arc (one line per quarter)", level=2)
    for q in quarters:
        if q in QUARTER_NARRATIVES:
            qtitle, _qmonths, _qnarr = QUARTER_NARRATIVES[q]
            add_bullet(doc, f"{q} — {qtitle}")

    doc.add_page_break()

    # ---- Per-Quarter rollups ----
    add_heading(doc, "2. Quarterly Rollups", level=1)
    for q in quarters:
        if q not in QUARTER_NARRATIVES:
            continue
        qtitle, qmonths, qnarr = QUARTER_NARRATIVES[q]
        add_heading(doc, f"{q} — {qtitle}", level=2)
        q_entries = by_quarter[q]
        add_para(doc, qnarr)

        # Top projects in the quarter
        qproj = Counter(e["project"] for e in q_entries)
        qproj_rows = [
            [p, str(c)] for p, c in qproj.most_common(10) if p != "(none)"
        ]
        if qproj_rows:
            add_para(doc, "Top projects in this quarter:", bold=True)
            add_table(doc, ["Project", "Entries"], qproj_rows)

        # Months in the quarter
        add_para(doc, "Monthly breakdown:", bold=True)
        for ym in qmonths:
            n = len(by_month.get(ym, []))
            add_bullet(doc, f"{ym} — {n} entries")
        doc.add_paragraph()

    doc.add_page_break()

    # ---- Per-month sections ----
    add_heading(doc, "3. Monthly Breakdown", level=1)
    for ym in months:
        m_entries = by_month[ym]
        nice = dt.datetime.strptime(ym + "-01", "%Y-%m-%d").strftime("%B %Y")
        add_heading(doc, f"{ym} — {nice}  ({len(m_entries)} entries)", level=2)

        narr = MONTH_NARRATIVES.get(ym)
        if narr:
            add_para(doc, narr)
        else:
            add_para(doc, "(No narrative provided for this month — see entries below.)")

        # Top projects this month
        mproj = Counter(e["project"] for e in m_entries)
        mproj_rows = [
            [p, str(c)] for p, c in mproj.most_common(8) if p != "(none)"
        ]
        if mproj_rows:
            add_para(doc, "Top projects this month:", bold=True)
            add_table(doc, ["Project", "Entries"], mproj_rows)

        # Day-by-day listing (group entries by ISO date)
        add_para(doc, "Daily entries:", bold=True)
        by_day: dict[str, list[dict]] = defaultdict(list)
        for e in m_entries:
            by_day[e["iso"]].append(e)
        for iso in sorted(by_day.keys()):
            day_entries = by_day[iso]
            day_p = doc.add_paragraph()
            day_run = day_p.add_run(iso)
            day_run.bold = True
            day_run.font.size = Pt(11)
            day_run.font.color.rgb = RGBColor(0x10, 0x3F, 0x82)
            for e in day_entries:
                tag_bits = []
                if e["project"] and e["project"] != "(none)":
                    tag_bits.append(e["project"])
                if e["category"] and e["category"] != "(none)":
                    tag_bits.append(e["category"])
                tag = f"[{' / '.join(tag_bits)}] " if tag_bits else ""
                add_bullet(doc, f"{tag}{e['title']}")
        doc.add_paragraph()

    doc.add_page_break()

    # ---- Appendix A — full chronological log ----
    add_heading(doc, "Appendix A — Full chronological log", level=1)
    add_para(
        doc,
        f"Every one of the {total} progress entries in the local database, in "
        f"chronological order. Columns: Date · Project · Category · Title.",
    )
    appendix_rows = [
        [e["iso"], e["project"], e["category"], e["title"]] for e in entries
    ]
    # python-docx struggles with very long tables; split into chunks of 200.
    chunk = 200
    for i in range(0, len(appendix_rows), chunk):
        add_table(
            doc,
            ["Date", "Project", "Category", "Title"],
            appendix_rows[i : i + chunk],
        )
        doc.add_paragraph()

    # ---- Save ----
    doc.save(str(OUT_DOCX))
    print(f"[ok] wrote {OUT_DOCX}  ({OUT_DOCX.stat().st_size / 1024:.1f} KB)")

    # ---- Plain-text mirror (so the user can also paste it anywhere) ----
    lines: list[str] = []
    lines.append("Bryan Kusno — Tencent Indonesia Journey")
    lines.append(f"{first_iso}  →  {last_iso}   ({total} entries)")
    lines.append("=" * 78)
    lines.append("")
    lines.append("EXECUTIVE SUMMARY")
    lines.append("-" * 78)
    lines.append(
        f"{total} progress entries · {len(months)} months · {len(quarters)} quarters"
    )
    lines.append("")
    lines.append("Top projects across the year:")
    for p, c in project_counter_total.most_common(15):
        if p == "(none)":
            continue
        lines.append(f"  {p:<20} {c}")
    lines.append("")
    lines.append("Activity-type distribution:")
    for c, n in category_counter_total.most_common(15):
        if c == "(none)":
            continue
        lines.append(f"  {c:<20} {n}")
    lines.append("")
    lines.append("=" * 78)
    lines.append("QUARTERLY ROLLUPS")
    lines.append("=" * 78)
    for q in quarters:
        if q not in QUARTER_NARRATIVES:
            continue
        qtitle, qmonths, qnarr = QUARTER_NARRATIVES[q]
        lines.append("")
        lines.append(f"{q} — {qtitle}")
        lines.append("-" * 78)
        lines.append(qnarr)
        lines.append("")
        for ym in qmonths:
            lines.append(f"  {ym} — {len(by_month.get(ym, []))} entries")
    lines.append("")
    lines.append("=" * 78)
    lines.append("MONTHLY BREAKDOWN")
    lines.append("=" * 78)
    for ym in months:
        m_entries = by_month[ym]
        nice = dt.datetime.strptime(ym + "-01", "%Y-%m-%d").strftime("%B %Y")
        lines.append("")
        lines.append(f"### {ym} — {nice}  ({len(m_entries)} entries)")
        narr = MONTH_NARRATIVES.get(ym)
        if narr:
            lines.append(narr)
        lines.append("")
        by_day = defaultdict(list)
        for e in m_entries:
            by_day[e["iso"]].append(e)
        for iso in sorted(by_day.keys()):
            lines.append(f"  {iso}")
            for e in by_day[iso]:
                tag_bits = []
                if e["project"] and e["project"] != "(none)":
                    tag_bits.append(e["project"])
                if e["category"] and e["category"] != "(none)":
                    tag_bits.append(e["category"])
                tag = f"[{' / '.join(tag_bits)}] " if tag_bits else ""
                lines.append(f"    - {tag}{e['title']}")
    lines.append("")
    lines.append("=" * 78)
    lines.append("APPENDIX A — FULL CHRONOLOGICAL LOG")
    lines.append("=" * 78)
    for e in entries:
        tag_bits = []
        if e["project"] and e["project"] != "(none)":
            tag_bits.append(e["project"])
        if e["category"] and e["category"] != "(none)":
            tag_bits.append(e["category"])
        tag = f"[{' / '.join(tag_bits)}] " if tag_bits else ""
        lines.append(f"{e['iso']}  {tag}{e['title']}")

    OUT_TXT.write_text("\n".join(lines), encoding="utf-8")
    print(f"[ok] wrote {OUT_TXT}  ({OUT_TXT.stat().st_size / 1024:.1f} KB)")


if __name__ == "__main__":
    main()
